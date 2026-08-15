from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DIAGRAMS = DOCS / "diagrams"
DIAGRAMS.mkdir(parents=True, exist_ok=True)


FUNCTIONAL_MERMAID = r'''flowchart LR
    U[门店老板 / 主管 / 员工] --> FE[Next.js 管理后台]
    FE --> AUTH[邮箱登录与权限]
    FE --> API[Django REST API]
    API --> DB[(PostgreSQL)]

    subgraph Master[基础资料]
      P[产品与配方]
      I[食材与库存单位]
      S[供应商管理]
      E[员工档案]
    end
    subgraph Operations[运营执行]
      PL[生产计划]
      R[采购入库 / 进货记录]
      SH[员工排班]
      EV[活动管理]
      SO[实际销售记录]
    end
    subgraph Analysis[经营分析]
      SA[销售分析]
      CM[成本管理]
      PA[盈利与产品表现]
    end

    API --> Master
    API --> Operations
    API --> Analysis
    P --> PL
    P --> SO
    S --> R
    I --> P
    I --> R
    PL -->|产品计划| DEMAND[食材需求计算]
    P -->|配方 BOM| DEMAND
    DEMAND --> INV[库存覆盖与采购预警]
    R --> INV
    SH -->|时薪 x 出勤时长| CM
    SO --> SA
    PL -->|实际制作| CM
    R -->|采购单价与入库数量| CM
    SA --> PA
    CM --> PA
    EV -.->|准备建议| PL
    EV -.->|库存建议| INV
'''

DATA_MERMAID = r'''erDiagram
    USER ||--o| USER_PREFERENCE : has
    USER }o--o{ ROLE : receives
    ROLE }o--o{ NAVIGATION_ITEM : grants
    PRODUCT ||--o{ RECIPE : has
    RECIPE ||--o{ RECIPE_SECTION : contains
    RECIPE_SECTION ||--o{ RECIPE_INGREDIENT : contains
    INGREDIENT ||--o{ RECIPE_INGREDIENT : used_by
    INGREDIENT ||--o| INVENTORY_ITEM : stocked_as
    SUPPLIER ||--o{ SUPPLIER_INGREDIENT : offers
    INGREDIENT ||--o{ SUPPLIER_INGREDIENT : offered_by
    INGREDIENT ||--o{ INVENTORY_RECEIPT : received
    SUPPLIER ||--o{ INVENTORY_RECEIPT : supplies
    PRODUCT ||--o{ PRODUCTION_PLAN : planned
    EMPLOYEE ||--o{ SCHEDULE_ENTRY : scheduled
    PRODUCT ||--o{ SALES_ORDER_LINE : sold
    SALES_ORDER ||--o{ SALES_ORDER_LINE : contains
    COST_ITEM ||--o{ MONTHLY_COST : records
    COST_MONTH ||--o{ MONTHLY_COST : groups
    EMPLOYEE ||--o{ SCHEDULE_ENTRY : earns
    BUSINESS_EVENT ||--o{ EVENT_CHECKLIST : has

    PRODUCT {
      uuid id PK
      string code UK
      string name_zh
      string name_en
      string sale_status
    }
    RECIPE {
      uuid id PK
      uuid product_id FK
      int yield_quantity
      string yield_unit
      bool is_active
    }
    INGREDIENT {
      uuid id PK
      string name UK
      string base_unit
      bool is_active
    }
    INVENTORY_ITEM {
      uuid id PK
      uuid ingredient_id FK
      decimal quantity
      decimal inventory_value
      int safety_buffer_days
    }
    PRODUCTION_PLAN {
      uuid id PK
      uuid product_id FK
      date planned_date
      int quantity
      int actual_quantity
      decimal actual_unit_material_cost
    }
    INVENTORY_RECEIPT {
      uuid id PK
      uuid ingredient_id FK
      uuid supplier_id FK
      decimal quantity
      decimal unit_price
      datetime received_at
    }
    SALES_ORDER {
      uuid id PK
      string reference UK
      datetime sold_at
    }
    SALES_ORDER_LINE {
      uuid id PK
      uuid order_id FK
      uuid product_id FK
      int quantity
      decimal paid_amount
      decimal refund_amount
    }
'''


def drawio_xml(
    title: str,
    cells: list[tuple[str, str, int, int, int, int]],
    edges: list[tuple[int, int, str]],
) -> str:
    nodes = [
        '<mxCell id="0"/>',
        '<mxCell id="1" parent="0"/>',
    ]
    for index, (label, style, x, y, width, height) in enumerate(cells, start=2):
        nodes.append(
            f'<mxCell id="{index}" value="{escape(label)}" style="{style}" vertex="1" parent="1">'
            f'<mxGeometry x="{x}" y="{y}" width="{width}" height="{height}" as="geometry"/>'
            '</mxCell>'
        )
    edge_start = len(cells) + 2
    edge_style = "edgeStyle=orthogonalEdgeStyle;rounded=1;orthogonalLoop=1;jettySize=auto;html=1;endArrow=block;endFill=1;"
    for offset, (source, target, label) in enumerate(edges):
        nodes.append(
            f'<mxCell id="{edge_start + offset}" value="{escape(label)}" style="{edge_style}" '
            f'edge="1" parent="1" source="{source + 2}" target="{target + 2}">'
            '<mxGeometry relative="1" as="geometry"/>'
            '</mxCell>'
        )
    return (
        '<mxfile host="app.diagrams.net" modified="2026-08-14T00:00:00.000Z" '
        'agent="BakeOps documentation generator" version="24.7.17">'
        f'<diagram id="bakeops" name="{escape(title)}"><mxGraphModel dx="1400" dy="900" grid="1" '
        'gridSize="10" guides="1" tooltips="1" connect="1" arrows="1" fold="1" page="1" pageScale="1" '
        'pageWidth="1600" pageHeight="1000"><root>'
        + "".join(nodes)
        + '</root></mxGraphModel></diagram></mxfile>'
    )


def write_diagrams() -> None:
    (DIAGRAMS / "bakeops-functional-architecture.mmd").write_text(FUNCTIONAL_MERMAID, encoding="utf-8")
    (DIAGRAMS / "bakeops-data-structure.mmd").write_text(DATA_MERMAID, encoding="utf-8")
    box = "rounded=1;whiteSpace=wrap;html=1;fillColor=#f5f7fa;strokeColor=#52606d;fontSize=14;"
    flow_cells = [
        ("门店用户", box, 80, 80, 160, 60),
        ("Next.js 管理后台", box, 320, 80, 190, 60),
        ("Django REST API", box, 590, 80, 180, 60),
        ("PostgreSQL", "shape=cylinder;whiteSpace=wrap;html=1;fillColor=#e8f1fb;strokeColor=#277da1;", 850, 80, 170, 70),
        ("产品与配方", box, 80, 230, 160, 60),
        ("供应商 / 进货", box, 300, 230, 180, 60),
        ("生产计划", box, 540, 230, 160, 60),
        ("库存管理", box, 760, 230, 160, 60),
        ("销售分析", box, 80, 380, 160, 60),
        ("成本管理", box, 300, 380, 160, 60),
        ("盈利与产品表现", box, 540, 380, 190, 60),
        ("活动管理", box, 790, 380, 160, 60),
    ]
    data_cells = [
        ("Product", box, 80, 80, 160, 60),
        ("Recipe / RecipeIngredient", box, 300, 80, 220, 60),
        ("Ingredient", box, 590, 80, 160, 60),
        ("InventoryItem", box, 820, 80, 170, 60),
        ("SupplierIngredient", box, 300, 220, 220, 60),
        ("InventoryReceipt", box, 590, 220, 180, 60),
        ("ProductionPlan", box, 80, 360, 180, 60),
        ("SalesOrder / Line", box, 330, 360, 190, 60),
        ("CostMonth / MonthlyCost", box, 590, 360, 220, 60),
        ("Employee / ScheduleEntry", box, 870, 360, 220, 60),
    ]
    flow_edges = [
        (0, 1, "使用"), (1, 2, "API"), (2, 3, "读写"),
        (4, 6, "产品"), (5, 7, "入库"), (6, 7, "需求"),
        (8, 10, "收入"), (9, 10, "成本"), (11, 6, "建议"),
    ]
    data_edges = [
        (0, 1, "拥有配方"), (1, 2, "使用食材"), (2, 3, "当前库存"),
        (4, 2, "供应食材"), (5, 2, "入库食材"), (6, 0, "计划产品"),
        (7, 0, "销售产品"), (8, 6, "物料成本"), (9, 8, "人工成本"),
    ]
    (DIAGRAMS / "bakeops-functional-architecture.drawio").write_text(
        drawio_xml("BakeOps 功能逻辑架构", flow_cells, flow_edges), encoding="utf-8"
    )
    (DIAGRAMS / "bakeops-data-structure.drawio").write_text(
        drawio_xml("BakeOps 数据结构", data_cells, data_edges), encoding="utf-8"
    )


def style_sheet(sheet) -> None:
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions
    for cell in sheet[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="277DA1")
        cell.alignment = Alignment(horizontal="center", vertical="center")
    for column in sheet.columns:
        width = min(max(max(len(str(cell.value or "")) for cell in column) + 2, 12), 42)
        sheet.column_dimensions[get_column_letter(column[0].column)].width = width
    for row in sheet.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)


def write_data_dictionary() -> None:
    workbook = Workbook()
    overview = workbook.active
    overview.title = "说明"
    overview.append(["BakeOps 数据字典", ""])
    overview.append(["版本", "2026-08-14"])
    overview.append(["用途", "描述当前数据库主要业务表、字段含义、关系和数据来源。"])
    overview.append(["主键约定", "业务模型默认使用 UUID；展示名称不是关联键。"])
    overview.append(["金额约定", "金额使用 Decimal，默认货币为 GBP。"])
    overview.append(["单位约定", "配方按投入单位计算；库存保留 base_unit 和展示单位；采购单价必须标明 price_unit。"])
    overview.append(["历史数据", "员工、产品和供应商等涉及历史引用的数据优先软删除或停用。"])
    overview.column_dimensions["A"].width = 18
    overview.column_dimensions["B"].width = 110
    for row in overview.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    overview["A1"].font = Font(bold=True, size=14, color="FFFFFF")
    overview["A1"].fill = PatternFill("solid", fgColor="277DA1")
    overview["B1"].fill = PatternFill("solid", fgColor="277DA1")

    tables = [
        ["模块", "表 / 模型", "用途", "核心关系", "当前状态"],
        ["用户与权限", "User", "登录账户、邮箱、用户名、姓名、状态", "User -> UserPreference / Role", "已实现"],
        ["用户与权限", "UserPreference", "语言、主题、时区、分页、侧边栏固定状态", "User 一对一", "已实现"],
        ["用户与权限", "Role / NavigationItem", "角色和页面权限", "Role 多对多 NavigationItem", "已实现"],
        ["产品", "Product", "销售产品主数据", "Product -> Recipe", "已实现"],
        ["产品", "Ingredient", "食材基础资料和 base_unit", "Ingredient -> RecipeIngredient / InventoryItem", "已实现"],
        ["产品", "Recipe / RecipeSection / RecipeIngredient", "产品配方 BOM、分区和产量", "Product -> Recipe -> items", "已实现"],
        ["供应商", "Supplier", "供应商主数据", "Supplier -> SupplierIngredient", "已实现"],
        ["供应商", "SupplierIngredient", "供应商食材报价和采购条件", "SupplierIngredient -> Supplier + Ingredient", "已实现"],
        ["库存", "InventoryItem", "当前库存数量和库存价值", "Ingredient 一对一", "已实现"],
        ["库存", "InventoryReceipt", "采购入库、单价、供应商和时间", "Ingredient + Supplier", "已实现"],
        ["库存", "ProductionPlan", "每日产品计划和实际制作", "Product + planned_date", "已实现"],
        ["人员", "Employee", "员工档案、入职/离职和时薪", "Employee -> ScheduleEntry", "已实现"],
        ["人员", "ScheduleEntry", "员工排班、工时和工资基础数据", "Employee + work_date", "已实现"],
        ["销售", "SalesOrder / SalesOrderLine", "真实订单、销售数量、折扣和退款", "SalesOrder -> lines -> Product", "已实现"],
        ["成本", "CostItem", "可选经营成本项目", "CostItem -> MonthlyCost", "已实现"],
        ["成本", "CostMonth / MonthlyCost", "每月实际成本和自动物料成本", "CostMonth -> MonthlyCost", "已实现"],
        ["活动", "Holiday / BusinessEvent / BusinessClosure", "节假日、活动和营业状态", "BusinessEvent -> Checklist", "已实现/持续增强"],
    ]
    table_sheet = workbook.create_sheet("业务表")
    for row in tables:
        table_sheet.append(row)
    style_sheet(table_sheet)

    fields = [
        ["表 / 模型", "字段", "类型", "是否必填", "说明", "业务规则 / 来源"],
        ["Product", "id", "UUID", "是", "产品主键", "不使用产品名称作为外键"],
        ["Product", "code", "string", "是", "产品编码", "唯一"],
        ["Product", "name_zh / name_en", "string", "是", "中英文产品名", "英文以中文业务含义为基准"],
        ["Recipe", "yield_quantity", "positive integer", "是", "一次配方产出多少可销售单位", "配方成本按产量折算为单份成本"],
        ["RecipeIngredient", "weight", "Decimal", "是", "一次配方投入量", "必须结合 unit 解读"],
        ["RecipeIngredient", "unit", "string", "是", "配方投入单位", "g/kg/ml/L/pcs 等"],
        ["Ingredient", "base_unit", "string", "是", "库存基础单位", "用于需求、库存和平均成本换算"],
        ["InventoryItem", "quantity", "Decimal", "是", "当前可用库存", "基础单位数量"],
        ["InventoryItem", "inventory_value", "Decimal", "否", "当前库存账面价值", "移动加权平均成本的分子"],
        ["InventoryReceipt", "quantity / unit", "Decimal / string", "是", "本次入库数量和单位", "真实采购入库记录"],
        ["InventoryReceipt", "unit_price / price_unit", "Decimal / string", "否", "实际采购单价及计价单位", "不能与供应商报价混淆"],
        ["ProductionPlan", "planned_date", "date", "是", "生产日期", "Product + date 唯一"],
        ["ProductionPlan", "quantity", "positive integer", "是", "计划制作数量", "未来计划使用"],
        ["ProductionPlan", "actual_quantity", "positive integer/null", "否", "实际制作数量", "过去日期的实际结果"],
        ["ProductionPlan", "actual_unit_material_cost", "Decimal/null", "否", "实际制作时捕获的单位材料成本", "历史盈利/物料成本计算使用"],
        ["Employee", "hire_date / departure_date", "date/null", "否", "入职和离职时间", "离职员工历史排班保留"],
        ["Employee", "hourly_rate", "Decimal", "是", "员工时薪", "排班工资 = 实际工时 x 时薪"],
        ["ScheduleEntry", "work_date", "date", "是", "排班日期", "当前工资模块按排班计算"],
        ["ScheduleEntry", "start_time / end_time / break_minutes", "time/time/int", "是", "班次时间和休息分钟", "计算实际工时"],
        ["SalesOrderLine", "quantity", "positive integer", "是", "真实售出数量", "不能超过实际生产数量的模拟约束"],
        ["SalesOrderLine", "standard_sales_amount", "Decimal", "是", "按标准售价的销售额", "用于分析折扣"],
        ["SalesOrderLine", "paid_amount / refund_amount", "Decimal", "是", "实际支付和退款", "净收入 = paid - refund"],
        ["MonthlyCost", "amount", "Decimal", "是", "当月实际成本金额", "员工工资、物料成本自动计算；其他成本手动维护"],
        ["MonthlyCost", "cost_month", "date", "是", "所属月份", "按 YYYY-MM 业务展示"],
        ["BusinessEvent", "start_date / end_date", "date", "是", "活动周期", "活动状态根据准备期自动计算"],
    ]
    field_sheet = workbook.create_sheet("字段字典")
    for row in fields:
        field_sheet.append(row)
    style_sheet(field_sheet)

    api_rows = [
        ["领域", "方法", "路径", "用途"],
        ["认证", "POST", "/api/v1/users/auth/login/", "邮箱登录"],
        ["导航", "GET", "/api/v1/navigation/menus/main-sidebar/tree/", "读取动态侧边栏"],
        ["产品", "GET/POST", "/api/v1/products/", "产品列表和新增"],
        ["配方", "POST/PUT/DELETE", "/api/v1/products/.../ingredients/", "维护配方原料"],
        ["供应商", "GET/POST", "/api/v1/suppliers/", "供应商列表和新增"],
        ["库存", "GET", "/api/v1/inventory/overview/", "库存需求和采购预警"],
        ["进货", "GET/POST", "/api/v1/inventory/receipts/", "采购入库和进货记录"],
        ["生产", "GET/POST/PUT", "/api/v1/inventory/production-plans/", "生产计划和实际制作"],
        ["排班", "GET/POST/PUT/DELETE", "/api/v1/schedules/", "员工排班"],
        ["销售分析", "GET", "/api/v1/sales/analysis/", "真实销售 KPI、趋势、产品和时段"],
        ["盈利分析", "GET", "/api/v1/sales/profitability/", "门店盈利和产品贡献毛利"],
        ["成本", "GET/PUT", "/api/v1/costs/overview/", "每月成本汇总"],
        ["活动", "GET/POST/PUT/DELETE", "/api/v1/events/", "活动、清单和营业状态"],
    ]
    api_sheet = workbook.create_sheet("API 目录")
    for row in api_rows:
        api_sheet.append(row)
    style_sheet(api_sheet)
    workbook.save(DOCS / "BakeOps_Data_Dictionary.xlsx")


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def write_owner_guide() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    title = document.add_heading("BakeOps 门店运营系统使用说明书", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph("面向餐厅老板与门店主管的简明版\n版本：2026-08-15")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(document, "一、BakeOps 是做什么的？")
    document.add_paragraph("BakeOps 把门店每天需要记录的事情集中到一个系统里：今天生产什么、库存还能用多久、从哪家供应商进货、员工排了什么班、实际卖了多少钱，以及最后到底赚了多少钱。")
    document.add_paragraph("系统的核心思路是：前面记录真实业务，后面自动计算结果。老板不需要手工把生产、库存、销售和成本再抄一遍。")
    document.add_paragraph("最重要的业务链：")
    document.add_paragraph("产品与配方 → 生产计划 → 食材需求 → 库存管理 → 采购入库 → 销售记录 → 成本管理 → 盈利分析")

    add_heading(document, "二、每天最常用的工作顺序")
    steps = [
        "先看生产计划：确认今天和未来几天要做哪些产品、每种做多少。",
        "再看库存管理：系统会根据生产计划和配方估算食材需求，并提示哪些食材需要采购。",
        "货到店后做采购入库：填写食材、数量、单价、采购时间和供应商。库存数量和成本会自动更新。",
        "员工班次完成后维护排班记录：工资按照实际排班时长和员工时薪计算。",
        "每天或每天营业结束后导入/记录实际销售：销售分析只看真实卖出和实际收款。",
        "月底查看成本管理和盈利分析：确认物料、工资、房租、水电等成本，以及最终经营利润。",
    ]
    for index, step in enumerate(steps, 1):
        document.add_paragraph(f"{index}. {step}")

    add_heading(document, "三、各个页面怎么用？")
    sections = [
        ("产品与配方", ["先维护产品名称、售价和配方。", "配方产量表示一次制作能得到多少个可销售单位。", "修改配方产量时，配方原料重量会按比例同步变化。", "当前预估成本来自当前库存加权平均成本；没有可用库存成本的食材会提示无法计算。"]),
        ("生产计划", ["按日期新增产品计划，过去日期还可以填写实际制作数量。", "未来日期只填写计划数量。", "历史数据用于分析计划准确率和实际物料成本。", "生产计划不要求每天制作全部产品，限定产品只在需要的活动或季节日期安排。"]),
        ("库存管理", ["当前库存、未来生产需求、可覆盖生产日和预计不足日期由系统计算。", "休息日没有需求，不会消耗库存。", "点击食材可查看需求来源、推荐供应商和建议采购量。", "采购入库后，库存数量、库存价值、平均成本和采购状态会联动更新。"]),
        ("供应商管理", ["可以按供应商查看，也可以按食材查看所有供应商。", "维护单价、单位、MOQ、提前预订天数和首选供应商。", "供应商报价是参考信息，真实入库单价以进货记录为准。"]),
        ("进货记录", ["每一条记录代表一次已经完成的采购入库。", "记录采购时间、食材、数量、成本单价、供应商和备注。", "这些记录会影响库存平均成本和产品当前预估成本。"]),
        ("员工与排班", ["员工资料维护姓名、职位、时薪、入职时间和离职时间。", "删除或离职员工后，历史排班仍然保留并显示历史员工信息。", "当前工资按排班工作时长乘员工时薪计算，成本管理中的工资不可直接编辑。"]),
        ("销售分析", ["只回答实际卖了什么、收了多少钱。", "净销售收入 = 实际支付金额 - 退款。", "系统同时展示折扣、退款、平均客单价、每日销售明细和产品实际成交价。"]),
        ("成本管理", ["员工工资来自排班，食材物料来自生产和配方，均为自动计算。", "房租、水电、保险、软件、维修等其他经营成本由用户按月份录入。", "历史月份可以单独修改，不会影响其他月份。"]),
        ("盈利分析", ["门店层面：净销售收入减去物料、工资和其他经营成本。", "产品层面：实际净销售收入减对应材料成本，得到贡献毛利。", "房租和人工暂时不强行分摊给产品，避免产生看似精确但不可靠的产品净利润。"]),
        ("活动管理", ["维护节假日、促销、KOL 合作、客户活动和休息/停业日期。", "活动页面提供生产和库存准备建议，但第一版不会自动修改生产计划或采购订单。"]),
    ]
    for heading, bullets in sections:
        add_heading(document, heading, 2)
        add_bullets(document, bullets)

    add_heading(document, "四、各页面的计算逻辑")

    calculation_sections = [
        (
            "首页仪表盘",
            [
                "右上角日期决定本页的销售、订单、计划生产、实际生产和最近 7 天趋势。最近 7 天包含所选日期及之前 6 天。",
                "净销售 = 实际支付金额 - 退款。订单数按不同订单去重，销量为订单明细数量合计。",
                "成本预估 = 物料成本 + 人工成本 + 经营成本分摊 + 当日专项成本。这是经营管理预估，不是精确会计成本，也不是当天采购金额。",
                "仪表盘物料成本优先使用实际制作数量；实际数量尚未录入时使用计划制作数量。单位材料成本来自实际制作时保存的成本快照，缺少快照时使用当前库存成本。",
                "人工成本 = 班次时长扣除休息时间后乘员工时薪。当前版本排班表尚未区分计划工时与打卡工时。",
                "经营成本分摊 = 当月非物料经营成本 / 当月计划营业日。活动等专项成本在单日活动全额计入，多日活动按持续天数平均分摊。",
            ],
        ),
        (
            "产品与配方",
            [
                "食材当前平均成本 = 当前库存价值 / 当前库存数量，统一按食材基础单位计算。",
                "整批配方预估成本 = Σ（配方食材投入量换算为基础单位 × 食材当前平均成本）。页面显示的是整批成本。",
                "单份产品材料成本 = 整批配方预估成本 / 配方产量。生产、成本和盈利计算使用单份成本。",
                "修改配方产量时，系统按比例同步放大或缩小所有食材重量，因此整批成本同步变化，单份成本原则上保持稳定。",
                "任何食材缺少库存数量、库存价值或可换算单位时，产品成本标记为无法完整计算，不把缺失成本当作 £0。",
            ],
        ),
        (
            "生产计划",
            [
                "差异 = 实际制作数量 - 计划制作数量。完成率 = 实际制作数量 / 计划制作数量 × 100%。",
                "未来日期显示已计划；过去日期有实际数量则显示已完成，没有实际数量则提示缺少实际数据；当天达到计划量显示已完成，否则显示进行中。",
                "实际制作录入时保存当时的单位材料成本快照，使后续采购价格变化不会改写历史生产成本。",
            ],
        ),
        (
            "库存管理",
            [
                "未来 14 天需求 = Σ（生产计划数量 / 配方产量 × 配方食材投入量）。当天只计算 max（计划数量 - 实际数量，0）的剩余需求。",
                "生产日日均需求 = 未来 14 天需求 / 其中实际有生产计划的天数；没有计划时显示无计划需求。",
                "系统按生产日期逐日扣减库存。可覆盖生产日是库存完整覆盖的计划生产日数量；第一次无法完整覆盖的日期为预计不足日期。",
                "建议采购日期 = 预计不足日期 - 首选供应商提前期 - 安全缓冲天数。",
                "紧急：按供应商提前期已无法在缺货前到货；需要采购：已经到达建议采购日期；关注：尚未到建议采购日期但计划窗口内会不足；正常：覆盖全部已制定计划；无计划需求：14 天内没有需求。",
                "建议采购量先补足 14 天计划缺口，再按照首选供应商 MOQ 的整数倍向上取整。",
            ],
        ),
        (
            "采购入库与进货记录",
            [
                "本次采购金额 = 入库数量换算到计价单位 × 实际采购单价。",
                "入库后库存数量 = 原库存数量 + 本次入库数量；库存价值 = 原库存价值 + 本次采购金额。",
                "新的库存平均成本 = 新库存总价值 / 新库存总数量。库存消耗按当时移动加权平均成本扣减库存价值。",
                "供应商报价只用于采购参考；进货单价记录真实历史交易；库存平均成本用于产品成本，三者不能混用。",
            ],
        ),
        (
            "员工排班与工资",
            [
                "可计薪工时 = 下班时间 - 上班时间 - 休息时间。每日工资 = 可计薪工时 × 员工时薪。",
                "月度员工工资 = 当月所有可计薪班次工资合计。工资在成本管理中只读，错误时应修改排班时间或员工时薪。",
                "员工离职或被删除后保留历史排班和当时显示的员工姓名，避免历史工资记录消失。",
            ],
        ),
        (
            "销售分析",
            [
                "净销售收入 = Σ（实际支付金额 - 退款金额）。标准销售额、折扣和退款分别保留，不能用销量乘当前标准售价代替真实收入。",
                "平均客单价 = 净销售收入 / 不同订单数。产品实际平均售价 = 产品净销售收入 / 产品销售数量。",
                "价格实现率 = 产品净销售收入 / 产品标准销售额 × 100%，用于观察折扣、赠送和退款对成交价格的影响。",
                "日、周、月只改变汇总粒度；开始日期和结束日期决定实际查询范围。",
            ],
        ),
        (
            "成本管理",
            [
                "员工工资自动来自排班，不能直接修改。食材物料成本自动来自生产记录和产品材料成本，也不能直接修改。",
                "历史日期物料成本 = 实际制作数量 × 当时保存的单位材料成本；当天成本由已完成实际数量和剩余计划数量拼接；未来日期使用计划制作数量和当前单位成本。",
                "月度物料成本 = 当月各生产日物料成本合计。房租、水电、保险、软件、维修等由用户按月份录入，每个月的数据独立保存。",
                "月度总成本 = 员工工资 + 食材物料成本 + 其他经营成本。饼图只展示金额大于 0 的当月项目。",
            ],
        ),
        (
            "盈利分析",
            [
                "毛利润 = 净销售收入 - 物料耗用成本。毛利率 = 毛利润 / 净销售收入 × 100%。",
                "经营利润 = 毛利润 - 人工成本 - 其他经营成本。经营利润率 = 经营利润 / 净销售收入 × 100%。",
                "趋势图中的净销售收入、毛利润和经营利润都使用同一日期范围；月度经营成本在趋势中按该月自然日分摊，月度总览仍使用整月实际金额。",
                "产品贡献毛利 = 产品实际净销售收入 - 产品对应材料成本。产品层不分摊房租、人工和水电，因此不称为产品净利润。",
                "产品四象限以净销售收入为横轴、贡献毛利率为纵轴，将产品分为明星、潜力、引流和待评估。",
            ],
        ),
        (
            "活动管理",
            [
                "准备期开始日 = 活动开始日 - 提前准备天数。系统比较准备期已经过去的比例与清单完成比例，明显落后时标记准备风险。",
                "建议生产数量 = 当前计划数量 ×（1 + 正向预计销售变化百分比）；第一版只给建议，不自动修改生产计划。",
                "活动额外食材需求 = 建议增加产品数量 / 配方产量 × 配方食材投入量，再与当前库存比较给出准备建议。",
                "节假日只作为参考；休息或停业日期默认不安排生产，但用户确认后可以覆盖。",
            ],
        ),
    ]
    for heading, bullets in calculation_sections:
        add_heading(document, heading, 2)
        add_bullets(document, bullets)

    add_heading(document, "五、老板最应该关注的三个数字")
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, ["指标", "它说明什么", "异常时先看哪里"]):
        cell.text = value
    rows = [
        ("净销售收入", "真实收了多少钱", "销售分析、折扣和退款"),
        ("经营利润", "扣除全部主要经营成本后还剩多少", "成本管理和销售分析"),
        ("产品贡献毛利", "哪个产品真正贡献了更多毛利", "盈利分析、产品配方和进货单价"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = value

    add_heading(document, "六、几个容易混淆的概念")
    add_bullets(document, [
        "供应商报价：供应商现在报给你的参考价格。",
        "进货单价：某一次实际采购入库的价格。",
        "库存平均成本：当前库存批次按照数量加权后的平均成本。",
        "产品当前预估成本：当前配方按当前库存平均成本计算的理论材料成本。",
        "产品贡献毛利：实际净销售收入减材料成本；不是把房租、工资也分摊后的最终净利润。",
    ])

    add_heading(document, "七、手机访问")
    document.add_paragraph("手机和电脑连接同一 Wi-Fi 后，在手机浏览器输入电脑的局域网地址，例如 http://192.168.1.61:3100。手机不能使用 localhost，因为 localhost 指的是手机自己。")

    add_heading(document, "八、数据安全和日常习惯")
    add_bullets(document, [
        "不要把真实客户、员工敏感信息和生产数据库直接提交到公开仓库。",
        "每天录入真实销售和采购入库，月底再补齐房租、水电等账单。",
        "重要业务数据修改前先确认日期范围，尤其是历史月份。",
        "如果成本或利润异常，按销售 → 生产 → 配方 → 进货 → 排班的顺序检查。",
    ])
    document.add_paragraph("技术人员请阅读项目根目录 README.md、docs/diagrams 和 BakeOps_Data_Dictionary.xlsx。")
    document.save(DOCS / "BakeOps_Owner_User_Guide.docx")


if __name__ == "__main__":
    write_diagrams()
    write_data_dictionary()
    write_owner_guide()
    print("Generated BakeOps diagrams, data dictionary and owner guide.")
