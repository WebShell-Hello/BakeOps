from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
ORIGINAL_SOURCE = Path("/Users/mankakiu/Desktop/Bakery_Operations_Website_Design_Spec_V1.0_CN.docx")
V11_SOURCE = PROJECT_ROOT / "docs" / "Bakery_Operations_Website_Design_Spec_V1.1_CN.docx"
SOURCE = ORIGINAL_SOURCE if ORIGINAL_SOURCE.exists() else V11_SOURCE
OUTPUT = PROJECT_ROOT / "docs" / "Bakery_Operations_Website_Design_Spec_V1.2_CN.docx"

ACCENT = "246B61"
ACCENT_LIGHT = "E8F3F0"
HEADER_TEXT = "FFFFFF"
GRID = "B8C8C4"
MUTED = RGBColor(88, 98, 96)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)


def style_table(table, widths: list[float] | None = None) -> None:
    table.style = "Table Grid"
    table.autofit = False
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            if widths and column_index < len(widths):
                cell.width = Cm(widths[column_index])
            if row_index == 0:
                set_cell_shading(cell, ACCENT)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(HEADER_TEXT)
            elif row_index % 2 == 0:
                set_cell_shading(cell, "F5F8F7")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.space_before = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is not None:
        for border in borders:
            border.set(qn("w:color"), GRID)


def insert_paragraph_before(anchor, text: str = "", style: str | None = None):
    paragraph = anchor.insert_paragraph_before(text, style=style)
    paragraph.paragraph_format.space_after = Pt(5)
    return paragraph


def insert_table_before(document: Document, anchor, rows: list[list[str]], widths: list[float] | None = None):
    table = document.add_table(rows=1, cols=len(rows[0]))
    for index, value in enumerate(rows[0]):
        set_cell_text(table.rows[0].cells[index], value, bold=True, color=HEADER_TEXT)
    for values in rows[1:]:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            set_cell_text(cells[index], value)
    style_table(table, widths)
    anchor._p.addprevious(table._tbl)
    return table


def add_bullets_before(anchor, items: list[str]) -> None:
    for item in items:
        insert_paragraph_before(anchor, item, "List Bullet")


def add_numbered_before(anchor, items: list[str]) -> None:
    for index, item in enumerate(items, start=1):
        paragraph = insert_paragraph_before(anchor, f"{index}. {item}")
        paragraph.paragraph_format.left_indent = Cm(0.65)
        paragraph.paragraph_format.first_line_indent = Cm(-0.45)


def find_paragraph(document: Document, exact: str):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == exact:
            return paragraph
    raise ValueError(f"Paragraph not found: {exact}")


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def append_directory_table(document: Document, anchor) -> None:
    rows = [
        ["路径", "职责"],
        ["frontend/src/app", "Next.js App Router 页面入口与路由。"],
        ["frontend/src/components", "按 auth、dashboard、users、access、navigation、employees、scheduling、products、ui 分域的界面组件。"],
        ["frontend/src/lib/api.ts", "前端唯一 API 访问层；页面不直接读取本地 JSON/CSV。"],
        ["backend/bakeops/<domain>", "Django 领域应用；通常包含 models、serializers、views、urls、permissions、tests 与 migrations。"],
        ["backend/config/settings", "base/local/test/production 分环境配置。"],
        ["docs/architecture", "架构决策记录（ADR）；记录长期有效的技术取舍。"],
        ["compose.yaml / .env.example", "本地容器编排与环境变量模板。"],
    ]
    insert_table_before(document, anchor, rows, [5.1, 11.7])


def find_table(document: Document, headers: tuple[str, ...]):
    for table in document.tables:
        if not table.rows:
            continue
        values = tuple(cell.text.strip() for cell in table.rows[0].cells)
        if values == headers:
            return table
    raise ValueError(f"Table not found: {headers}")


def upgrade_v11_to_v12(document: Document) -> None:
    replace_paragraph_text(
        find_paragraph(document, "版本：V1.1（规划基线 + 实施状态）"),
        "版本：V1.2（规划基线 + 实施状态）",
    )
    replace_paragraph_text(
        find_paragraph(document, "4.4 Purchases & Suppliers｜采购与供应商"),
        "4.4 Supplier Management｜供应商管理",
    )
    replace_paragraph_text(
        find_paragraph(document, "页面形态：列表。记录供应来源、采购价格、采购订单和到货状态，为成本和库存提供来源。"),
        "页面形态：供应商列表 + 详情 Drawer。维护食材从哪里采购，以及每家供应商的价格、MOQ、提前预订和首选关系；采购订单与到货状态在后续独立模块实现。",
    )
    summary_paragraphs = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip() == "顶部摘要/系统计算："
    ]
    replace_paragraph_text(summary_paragraphs[3], "列表重点信息：")
    replace_paragraph_text(find_paragraph(document, "本月采购额（可选）"), "供应食材数")
    replace_paragraph_text(find_paragraph(document, "价格变动提示（后续）"), "地址、联系方式与备注")

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip() == "Purchases & Suppliers":
                    cell.text = "Supplier Management"
                elif cell.text.strip() == "维护采购与供应商":
                    cell.text = "维护供应商与食材采购条件"
            if len(row.cells) >= 3 and row.cells[1].text.strip() == "Supplier Management":
                if row.cells[0].text.strip() == "P1":
                    row.cells[2].text = "已实现供应商与食材采购条件；采购单与到货后续建设"

    supplier_fields = document.tables[13]
    supplier_field_names = [
        "供应商名称",
        "地址",
        "联系人",
        "电话",
        "Email",
        "备注",
        "食材",
        "单价与计价单位",
        "MOQ 与单位",
        "提前预订天数",
        "采购条件备注",
        "供应状态",
        "首选供应商",
    ]
    while len(supplier_fields.rows) < len(supplier_field_names) + 1:
        supplier_fields.add_row()
    for index, field_name in enumerate(supplier_field_names, start=1):
        supplier_fields.rows[index].cells[0].text = field_name
        supplier_fields.rows[index].cells[1].text = ""
    document.tables[14].cell(0, 0).text = (
        "关联逻辑  Supplier → SupplierIngredient ← Ingredient。一个供应商可供应多种食材；"
        "一种食材可对应多个供应商，但同一时间最多只有一个启用中的首选供应商。"
    )
    for row in document.tables[35].rows:
        if row.cells[0].text.strip() == "Supplier":
            row.cells[1].text = "定义供应商基础资料"
            row.cells[2].text = "通过 SupplierIngredient 关联 Ingredient、价格、MOQ、提前天数与首选状态"

    implemented = find_table(document, ("模块", "状态", "当前已实现能力"))
    cells = implemented.add_row().cells
    for cell, value in zip(
        cells,
        (
            "供应商管理",
            "已实现",
            "供应商基础资料、搜索和详情 Drawer；按当前配方食材维护价格、单位、MOQ、提前天数、备注、停用和首选供应商；支持一种食材对应多个供应商。",
        ),
        strict=True,
    ):
        set_cell_text(cell, value)

    pending = find_table(document, ("模块", "状态", "当前边界 / 下一步"))
    for row in pending.rows:
        if row.cells[0].text.strip() == "原料成本":
            row.cells[0].text = "采购与原料成本"
            row.cells[2].text = "供应商报价与采购条件已经建立；价格历史、采购单、入库计价和配方 estimated_price 自动计算尚未完成。"
        elif row.cells[0].text.strip() == "Sales / Production / Inventory / Purchase / Waste":
            row.cells[0].text = "Sales / Production / Inventory / Purchase Orders / Waste"

    domains = find_table(document, ("后端应用", "主要实体 / 职责"))
    api_row = next(row for row in domains.rows if row.cells[0].text.strip() == "api")
    cells = domains.add_row().cells
    set_cell_text(cells[0], "suppliers")
    set_cell_text(cells[1], "Supplier、SupplierIngredient、食材报价与采购条件。")
    api_row._tr.addnext(cells[0]._tc.getparent())

    milestones = find_table(document, ("里程碑", "状态", "交付内容"))
    for row in milestones.rows:
        if row.cells[0].text.strip() == "M3 库存与采购基础":
            row.cells[1].text = "进行中"
            row.cells[2].text = "供应商及食材采购条件已完成；下一步建设报价历史、采购单、库存批次/流水、安全库存和临期状态。"

    next_steps = {
        "1. 完成原料库、供应商和采购价格模型，给现有配方 estimated_price 提供可信来源。": "1. 在现有供应商采购条件之上增加报价历史和采购单，避免直接覆盖历史价格。",
        "2. 建立库存批次与 Inventory Movement，不把当前库存仅保存为可随意覆盖的数字。": "2. 完善采购单位与配方单位换算，并将有效采购价格接入配方 estimated_price。",
        "3. 实现生产计划：计划产品数量 × 当前有效配方，汇总原料需求并比较可用库存。": "3. 建立库存批次与 Inventory Movement，不把当前库存仅保存为可随意覆盖的数字。",
        "4. 实现实际生产与损耗，让计划、生产、销售、损耗形成产品去向核算。": "4. 实现生产计划：计划产品数量 × 当前有效配方，汇总原料需求并比较可用库存。",
        "5. 接入销售记录后统一成本与利润口径，再把 Dashboard 空状态替换为真实指标。": "5. 实现实际生产与损耗，让计划、生产、销售、损耗形成产品去向核算。",
        "6. 最后推进日历活动、营销、分析和需求预测，避免在基础事实数据不足时过早做智能建议。": "6. 接入销售记录后统一成本与利润口径，再把 Dashboard 空状态替换为真实指标。",
    }
    for old, new in next_steps.items():
        replace_paragraph_text(find_paragraph(document, old), new)

    for run in document.sections[0].footer.paragraphs[0].runs:
        if "V1.1" in run.text:
            run.text = run.text.replace("V1.1", "V1.2")

    document.core_properties.subject = "业务蓝图、当前实施状态、技术架构、工程约定与路线图"
    document.save(OUTPUT)


def main() -> None:
    document = Document(SOURCE)
    if SOURCE == V11_SOURCE:
        upgrade_v11_to_v12(document)
        print(OUTPUT)
        return
    document.core_properties.title = "Bakery Operations & Management Platform - 网站设计说明书"
    document.core_properties.subject = "业务蓝图、当前实施状态、技术架构、工程约定与路线图"

    replace_paragraph_text(find_paragraph(document, "版本：V1.0"), "版本：V1.2（规划基线 + 实施状态）")
    replace_paragraph_text(find_paragraph(document, "日期：2026年8月13日"), "更新日期：2026年8月14日")

    positioning = document.tables[0].cell(0, 0)
    positioning.text = (
        "文档定位\n"
        "本说明书既是业务与产品蓝图，也是当前工程实施的权威上下文。第 1—12 章保留目标能力与长期设计；"
        "第 13—18 章记录截至 2026 年 8 月 14 日的真实技术架构、工程约定、已实现范围、当前边界和后续路线。"
        "后续 Agent、开发者或大模型应先核对代码与测试，再更新本文件中的实施状态，不能把规划项直接视为已完成。"
    )
    set_cell_shading(positioning, ACCENT_LIGHT)
    for index, run in enumerate(positioning.paragraphs[0].runs):
        run.font.color.rgb = RGBColor.from_string(ACCENT)
        run.bold = index == 0

    toc = document.tables[1]
    for chapter, title in [
        ("13", "当前项目与实施状态（2026-08-14）"),
        ("14", "技术架构与运行方式"),
        ("15", "工程与产品约定"),
        ("16", "里程碑、当前进度与下一步"),
        ("17", "Agent / 大模型接手指南"),
        ("18", "未来展望"),
    ]:
        cells = toc.add_row().cells
        cells[0].text = chapter
        cells[1].text = title

    replace_paragraph_text(
        find_paragraph(document, "第一阶段建议至少设计三类角色。即使 V1.0 暂时不实现完整 RBAC，也应在页面和数据模型上预留权限边界。"),
        "系统已实现可配置的角色与页面级访问控制。Owner / Admin、Store Manager、Staff 仍是业务角色基线，实际角色由管理员在“角色权限”页面维护；用户可拥有多个角色，最终页面权限取各角色权限并集。",
    )
    replace_paragraph_text(
        find_paragraph(document, "左侧导航建议按“总览—日常运营—产品—计划—人员—分析”组织，避免把数据库实体直接等同为导航菜单。原材料虽然是重要基础数据，但不需要单独占用一级导航。"),
        "左侧导航按“总览—日常运营—产品—计划—人员—分析—系统设置”组织。导航已由数据库配置驱动，支持中英文标签、图标、路径、显示状态和拖拽排序；用户只看到其角色允许访问的页面。原材料仍作为产品与配方内部基础库，不单独占用一级导航。",
    )
    replace_paragraph_text(
        find_paragraph(document, "4.4 Purchases & Suppliers｜采购与供应商"),
        "4.4 Supplier Management｜供应商管理",
    )
    replace_paragraph_text(
        find_paragraph(document, "页面形态：列表。记录供应来源、采购价格、采购订单和到货状态，为成本和库存提供来源。"),
        "页面形态：供应商列表 + 详情 Drawer。维护食材从哪里采购，以及每家供应商的价格、MOQ、提前预订和首选关系；采购订单与到货状态在后续独立模块实现。",
    )
    summary_paragraphs = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip() == "顶部摘要/系统计算："
    ]
    replace_paragraph_text(summary_paragraphs[3], "列表重点信息：")
    replace_paragraph_text(find_paragraph(document, "本月采购额（可选）"), "供应食材数")
    replace_paragraph_text(find_paragraph(document, "价格变动提示（后续）"), "地址、联系方式与备注")

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip() == "Purchases & Suppliers":
                    cell.text = "Supplier Management"
                elif cell.text.strip() == "维护采购与供应商":
                    cell.text = "维护供应商与食材采购条件"
        for row in table.rows:
            if len(row.cells) >= 3 and row.cells[1].text.strip() == "Supplier Management":
                if row.cells[0].text.strip() == "P1":
                    row.cells[2].text = "已实现供应商与食材采购条件；采购单与到货后续建设"

    master_data_table = document.tables[35]
    for row in master_data_table.rows:
        if row.cells[0].text.strip() == "Supplier":
            row.cells[1].text = "定义供应商基础资料"
            row.cells[2].text = "通过 SupplierIngredient 关联 Ingredient、价格、MOQ、提前天数与首选状态"

    supplier_fields = document.tables[13]
    supplier_field_names = [
        "供应商名称",
        "地址",
        "联系人",
        "电话",
        "Email",
        "备注",
        "食材",
        "单价与计价单位",
        "MOQ 与单位",
        "提前预订天数",
        "采购条件备注",
        "供应状态",
        "首选供应商",
    ]
    while len(supplier_fields.rows) < len(supplier_field_names) + 1:
        supplier_fields.add_row()
    for index, field_name in enumerate(supplier_field_names, start=1):
        supplier_fields.rows[index].cells[0].text = field_name
        supplier_fields.rows[index].cells[1].text = ""

    supplier_logic = document.tables[14].cell(0, 0)
    supplier_logic.text = (
        "关联逻辑  Supplier → SupplierIngredient ← Ingredient。一个供应商可供应多种食材；"
        "一种食材可对应多个供应商，但同一时间最多只有一个启用中的首选供应商。"
    )

    product_fields = document.tables[17]
    additions = [
        ("配方产量", "一个配方批次可产出的产品数量及单位；修改产量时按比例缩放每项原料重量。"),
        ("配方总重量", "由当前有效配方中所有原料重量自动汇总。"),
    ]
    for label, description in additions:
        cells = product_fields.add_row().cells
        cells[0].text = label
        cells[1].text = description

    replace_paragraph_text(
        find_paragraph(document, "角色权限与审计日志：老板、店长、员工看到不同数据并保留修改记录。"),
        "权限深化与审计日志：角色及页面访问控制已经实现；后续增加字段级/动作级权限与完整操作审计记录。",
    )

    footer = document.sections[0].footer.paragraphs[0]
    for run in footer.runs:
        if "V1.0" in run.text:
            run.text = run.text.replace("V1.0", "V1.2")

    appendix_anchor = find_paragraph(document, "附录 A｜页面形态总表")
    page_break = insert_paragraph_before(appendix_anchor)
    page_break.add_run().add_break(WD_BREAK.PAGE)

    insert_paragraph_before(appendix_anchor, "13. 当前项目与实施状态（2026-08-14）", "Heading 1")
    insert_paragraph_before(
        appendix_anchor,
        "BakeOps 当前是一套面向单店面包店的生产级运营管理平台基础版本。项目正在先建立可靠的身份、权限、导航和主数据能力，再依照业务闭环逐步接入库存、采购、生产、销售、损耗、成本和分析。当前开发数据为合成数据，但系统架构、数据库、API、权限和迁移均采用真实生产形态。",
    )

    insert_paragraph_before(appendix_anchor, "13.1 状态定义", "Heading 2")
    insert_table_before(
        document,
        appendix_anchor,
        [
            ["状态", "判定标准"],
            ["已实现", "存在真实数据库模型/API/页面，主要工作流可操作，并有相应自动化测试或实际联调验证。"],
            ["部分实现", "已有页面骨架或部分业务能力，但关键数据闭环、计算、权限或操作仍未完成。"],
            ["待开发", "仍属于规划蓝图，当前代码中没有可用业务实现。"],
        ],
        [3.2, 13.6],
    )

    insert_paragraph_before(appendix_anchor, "13.2 已实现功能模块", "Heading 2")
    implemented_rows = [
        ["模块", "状态", "当前已实现能力"],
        ["应用骨架与响应式布局", "已实现", "登录后管理后台、顶部栏、可折叠/移动端侧边栏、面包屑、统一空状态与消息提示；同一 Wi-Fi 下支持手机访问和登录。"],
        ["登录、注册与会话", "已实现", "邮箱登录；邮箱不区分大小写且唯一；用户名可重复、支持大小写；Session + CSRF；记住登录；退出；注册数字图形验证码。"],
        ["个人资料与偏好", "已实现", "用户可修改用户名、名、姓和密码；中英文、主题、时区、日期格式、周起始日、表格页大小等偏好可持久化。"],
        ["用户管理", "已实现", "搜索、新增、编辑、批量删除、锁定/解锁、重置密码、保护账户、多角色分配；超级管理员和当前账户有后端保护规则。"],
        ["角色与页面权限", "已实现", "角色 CRUD、软删除/恢复/永久删除、保护角色、按页面授权；用户多角色权限取并集。"],
        ["动态菜单管理", "已实现", "数据库驱动的分类与页面；中英文标签、图标、路径、显示/启用状态、拖拽排序；修订号防止并发覆盖。"],
        ["员工档案", "已实现", "员工号、姓名、性别、出生日期、职位、时薪、全/兼职、邮箱、状态；搜索、增改、软删除、恢复和批量操作。"],
        ["排班", "已实现", "基于员工的日期、班次、开始/结束时间、备注管理；日期范围读取、增改删、批量删除；离职/删除员工不会破坏历史排班。"],
        ["产品与配方", "已实现", "产品中英文名称、在售状态、备注、制作步骤；配方分区与原料 CRUD；配方总重量；配方产量及单位；修改产量时按比例缩放全部原料重量。"],
        ["供应商管理", "已实现", "供应商基础资料、搜索和详情 Drawer；按当前配方食材维护价格、单位、MOQ、提前天数、备注、停用和首选供应商；支持一种食材对应多个供应商。"],
        ["通用分页", "已实现", "左侧显示 Viewing x out of XX / 当前显示 x 条，共 XX 条；右侧 Rows per page / 每页行数采用数字输入框原生上下按钮，后接 Previous、当前页、Next。"],
        ["健康检查与 API 文档", "已实现", "前端展示 Next.js、Django API、PostgreSQL 状态；后端提供 /api/v1/health/、OpenAPI Schema 与 Swagger 文档。"],
    ]
    insert_table_before(document, appendix_anchor, implemented_rows, [4.0, 2.2, 10.6])

    insert_paragraph_before(appendix_anchor, "13.3 部分实现与待开发范围", "Heading 2")
    pending_rows = [
        ["模块", "状态", "当前边界 / 下一步"],
        ["Dashboard", "部分实现", "布局、KPI 卡片、系统状态和空状态已完成；尚未接入销售、库存、生产、损耗和利润真实指标。"],
        ["Attendance 实际考勤", "部分实现", "当前完成的是计划排班，不是签到、签退、休息时长、实际工时和人工成本核算。"],
        ["采购与原料成本", "部分实现", "供应商报价与采购条件已经建立；价格历史、采购单、入库计价和配方 estimated_price 自动计算尚未完成。"],
        ["Sales / Production / Inventory / Purchase Orders / Waste", "待开发", "导航与占位页存在，但尚无完整业务模型、API 和可用工作流。"],
        ["Production Planning / Events / Marketing", "待开发", "属于后续运营计划阶段。"],
        ["Analytics / Forecasting", "待开发", "需要先形成可靠的交易、库存、成本和历史数据。"],
        ["通知与全局搜索", "待开发", "顶部入口已展示，但当前没有可用业务逻辑。"],
        ["多门店、审计日志、生产部署", "待开发", "数据模型目前按单店起步；正式上线前还需安全加固、备份、监控、CI/CD 和部署方案。"],
    ]
    insert_table_before(document, appendix_anchor, pending_rows, [4.4, 2.2, 10.2])

    insert_paragraph_before(appendix_anchor, "14. 技术架构与运行方式", "Heading 1")
    insert_paragraph_before(appendix_anchor, "14.1 技术栈", "Heading 2")
    stack_rows = [
        ["层级", "技术与版本", "职责"],
        ["前端", "Next.js 16.3、React 19.2、TypeScript 5.9、Tailwind CSS 4", "App Router 管理界面、响应式交互、中英文与主题、统一 API 调用。"],
        ["前端组件", "Lucide React、TanStack Table、date-fns、React Day Picker、Recharts", "图标、表格、日期与未来图表能力；优先复用现有 UI 组件。"],
        ["后端", "Python 3.13、Django 5.2 LTS、Django REST Framework 3.16", "领域模型、REST API、Session 认证、权限、校验和业务规则。"],
        ["数据库", "PostgreSQL 17", "真实关系模型、UUID 主键、迁移、约束和索引。测试环境可使用 SQLite 内存库。"],
        ["接口文档", "drf-spectacular / OpenAPI", "Schema 与 Swagger API 文档。"],
        ["本地运行", "Docker Compose", "frontend、backend、postgres 三服务；持久卷和网络使用 BO- 前缀。"],
        ["质量工具", "ESLint、TypeScript、pytest、Ruff、mypy", "前端静态检查与后端测试、格式/类型检查。"],
    ]
    insert_table_before(document, appendix_anchor, stack_rows, [3.0, 5.8, 8.0])

    insert_paragraph_before(appendix_anchor, "14.2 请求与数据流", "Heading 2")
    add_numbered_before(
        appendix_anchor,
        [
            "浏览器访问 Next.js；前端业务代码统一通过 frontend/src/lib/api.ts 请求 /api/v1。",
            "Next.js 将同源 /api/v1 请求代理到 Docker 网络中的 Django backend，避免手机或跨端环境中的 Cookie/CSRF 跨源问题。",
            "Django REST Framework 执行 Session 认证、CSRF、角色权限和序列化校验，并访问 PostgreSQL。",
            "所有环境使用同一套模型、迁移、API 契约和权限逻辑；开发与演示只替换业务数据，不替换架构。",
        ],
    )

    insert_paragraph_before(appendix_anchor, "14.3 核心领域应用", "Heading 2")
    domain_rows = [
        ["后端应用", "主要实体 / 职责"],
        ["common", "BaseModel：UUID id、created_at、updated_at。"],
        ["users", "User、UserPreference、邮箱认证、注册验证码、会话、个人资料和用户管理。"],
        ["access", "Role 与页面权限关系。"],
        ["navigation", "NavigationMenu、NavigationItem、动态树、排序与可见性。"],
        ["employees", "Employee 员工主数据和软删除。"],
        ["scheduling", "ScheduleEntry 与员工排班。"],
        ["products", "Product、Ingredient、Recipe、RecipeSection、RecipeIngredient。"],
        ["suppliers", "Supplier、SupplierIngredient、食材报价与采购条件。"],
        ["api", "健康检查与各领域 API 汇总入口。"],
    ]
    insert_table_before(document, appendix_anchor, domain_rows, [4.2, 12.6])

    insert_paragraph_before(appendix_anchor, "14.4 目录职责", "Heading 2")
    append_directory_table(document, appendix_anchor)

    insert_paragraph_before(appendix_anchor, "14.5 本地开发运行", "Heading 2")
    add_bullets_before(
        appendix_anchor,
        [
            "启动：复制 .env.example 为 .env，执行 docker compose up --build -d。",
            "默认入口：前端 http://localhost:3100；API http://localhost:8000/api/v1；Django Admin http://localhost:8000/admin。",
            "同一 Wi-Fi 手机访问：使用电脑局域网 IP 加前端端口；该 IP 变化时同步更新 Django allowed hosts、CORS/CSRF trusted origins 与 Next allowedDevOrigins。",
            "常用命令：make up、make down、make logs、make check、make test、make migrate。",
            "环境变量中的密钥、数据库密码和默认密码仅为开发默认值，生产环境必须替换并通过安全配置注入。",
        ],
    )

    insert_paragraph_before(appendix_anchor, "15. 工程与产品约定", "Heading 1")
    conventions = [
        ["主题", "当前约定"],
        ["事实来源", "运行行为以代码、迁移和测试为准；本文件记录意图与快照。两者不一致时先验证代码，再同步文档。"],
        ["数据策略", "使用真实 PostgreSQL、REST API 和关系模型；开发/演示使用可重复的合成数据，禁止前端增加独立 JSON/CSV mock 路径。"],
        ["API", "业务接口统一使用 /api/v1；前端统一经 api.ts；保持请求/响应强类型和向后兼容。"],
        ["认证身份", "邮箱是登录标识和不区分大小写的唯一约束；用户名是显示属性，可重复并保留大小写。"],
        ["安全", "浏览器会话使用 Session + CSRF；敏感规则必须在后端执行，不能只依赖按钮隐藏。"],
        ["权限", "用户可分配多个角色，页面权限取并集；角色只关联 PAGE，不直接关联 CATEGORY；无可见子页面时分类自动隐藏。"],
        ["数据标识", "领域实体默认使用 UUID；关联使用稳定 ID，不使用显示名称作为外键。"],
        ["删除策略", "员工和角色等需要恢复历史的数据采用软删除；用户删除目前为真实删除，并保护超级管理员、当前账户和受保护账户。"],
        ["金额与数量", "金额使用 Decimal；数量必须有明确单位。配方重量使用 Decimal；产量变化按比例缩放配方项。"],
        ["国际化", "当前支持 zh-CN 与 en-GB；新增用户可见文案必须同时提供中英文，英文含义以已确认的中文版业务含义为基准。"],
        ["页面与组件", "沿用现有 DashboardShell、Button、Card、DataPagination、Toast、PageBreadcrumb 和 API 层，不重复实现同类组件。"],
        ["分页", "默认每页 10 条；数字输入支持键盘输入和浏览器原生上下按钮，不额外放置左右加减按钮；中英文文案保持一致。"],
        ["迁移", "任何持久模型变化都必须提交 Django migration；不得通过手工数据库改表代替迁移。"],
        ["测试", "业务规则和权限边界优先增加后端 API 测试；前端至少通过 lint 与 typecheck，关键页面需进行桌面与手机实际联调。"],
        ["变更范围", "遵循现有领域边界，优先小而完整的改动；不顺手重构无关代码，不覆盖其他开发者未提交修改。"],
    ]
    insert_table_before(document, appendix_anchor, conventions, [3.2, 13.6])

    insert_paragraph_before(appendix_anchor, "16. 里程碑、当前进度与下一步", "Heading 1")
    milestone_rows = [
        ["里程碑", "状态", "交付内容"],
        ["M0 工程基础", "完成", "Next.js + Django REST + PostgreSQL + Docker Compose；健康检查；真实 API/迁移/合成数据策略。"],
        ["M1 身份与治理", "完成", "邮箱认证、注册验证码、个人资料/偏好、用户管理、角色权限、动态菜单、中英文与主题。"],
        ["M2 核心主数据", "完成", "员工档案、计划排班、产品与配方、配方产量联动、通用分页。"],
        ["M3 库存与采购基础", "进行中", "供应商及食材采购条件已完成；下一步建设报价历史、采购单、库存批次/流水、安全库存和临期状态。"],
        ["M4 生产闭环", "计划中", "生产计划、配方耗材汇总、库存缺口、实际生产与异常记录。"],
        ["M5 销售与损耗", "计划中", "销售交易/导入、标准损耗原因、产品去向与库存扣减。"],
        ["M6 成本与 Dashboard", "计划中", "原料成本、人工成本、毛利/经营利润口径和真实 KPI/图表。"],
        ["M7 事件、营销与分析", "计划中", "活动日历、提醒、营销记录、分析页和有数据基础后的预测建议。"],
    ]
    insert_table_before(document, appendix_anchor, milestone_rows, [4.0, 2.3, 10.5])

    insert_paragraph_before(appendix_anchor, "16.1 推荐的近期推进顺序", "Heading 2")
    add_numbered_before(
        appendix_anchor,
        [
            "在现有供应商采购条件之上增加报价历史和采购单，避免直接覆盖历史价格。",
            "完善采购单位与配方单位换算，并将有效采购价格接入配方 estimated_price。",
            "建立库存批次与 Inventory Movement，不把当前库存仅保存为可随意覆盖的数字。",
            "实现生产计划：计划产品数量 × 当前有效配方，汇总原料需求并比较可用库存。",
            "实现实际生产与损耗，让计划、生产、销售、损耗形成产品去向核算。",
            "接入销售记录后统一成本与利润口径，再把 Dashboard 空状态替换为真实指标。",
        ],
    )

    insert_paragraph_before(appendix_anchor, "16.2 当前主要技术债与上线前事项", "Heading 2")
    add_bullets_before(
        appendix_anchor,
        [
            "生产环境密钥、默认密码、HTTPS、Secure Cookie、域名和反向代理配置尚需正式化。",
            "需要建立 CI：前端 lint/typecheck/build，后端 pytest/Ruff/mypy 和迁移检查。",
            "需要补充数据库备份恢复、日志、监控、错误追踪和容量策略。",
            "当前权限以页面访问和领域 manage_* 权限为主；财务敏感字段与动作级授权需在后续深化。",
            "实际考勤与计划排班必须保持为两个概念，后续不要直接用计划班次计算已发生工资。",
            "多门店尚未落库；引入 Store / Location 时应评估所有交易和主数据的归属关系。",
        ],
    )

    insert_paragraph_before(appendix_anchor, "17. Agent / 大模型接手指南", "Heading 1")
    insert_paragraph_before(
        appendix_anchor,
        "本节用于让后续 Agent、开发者或大模型在较短时间内理解项目，并降低基于旧文档或页面占位内容作出错误判断的风险。",
    )
    insert_paragraph_before(appendix_anchor, "17.1 建议阅读顺序", "Heading 2")
    add_numbered_before(
        appendix_anchor,
        [
            "阅读本文件第 13—18 章，理解当前状态、技术边界和下一步。",
            "阅读 README.md、compose.yaml、.env.example 和 docs/architecture 下的 ADR。",
            "通过 rg 搜索对应领域，先查看 models.py、serializers.py、views.py、urls.py、permissions.py 和 tests。",
            "查看 frontend/src/lib/api.ts 及对应 components/app 页面，确认前后端契约和现有交互。",
            "检查当前工作区修改，绝不回滚或覆盖无法确认来源的改动。",
            "实现后运行与风险匹配的检查，并在桌面与手机端验证关键用户流程。",
            "若完成度或架构发生变化，同步更新本说明书的状态日期、矩阵、里程碑与约定。",
        ],
    )

    insert_paragraph_before(appendix_anchor, "17.2 决策优先级", "Heading 2")
    priority_rows = [
        ["优先级", "依据"],
        ["1", "用户在当前任务中的明确要求。"],
        ["2", "现有可执行测试、数据库约束、API 契约和实际代码行为。"],
        ["3", "本文件第 13—18 章的当前实施约定。"],
        ["4", "第 1—12 章的长期业务蓝图。"],
        ["5", "合理推断；推断必须明确标注并尽快用代码或用户确认验证。"],
    ]
    insert_table_before(document, appendix_anchor, priority_rows, [2.3, 14.5])

    insert_paragraph_before(appendix_anchor, "17.3 完成定义（Definition of Done）", "Heading 2")
    add_bullets_before(
        appendix_anchor,
        [
            "功能连接真实 API 与数据库，不以只存在于浏览器内存的演示状态冒充完成。",
            "后端完成权限、校验、数据库约束和必要迁移；错误响应不会泄露敏感信息。",
            "前端具备加载、空数据、错误、提交中、成功反馈和响应式状态。",
            "所有新增用户可见文字同时提供中文和英文。",
            "高风险业务规则有自动化测试；前端通过 lint/typecheck，关键流程完成实际联调。",
            "文档、API 类型与实现同步；明确记录仍未完成的边界。",
        ],
    )

    insert_paragraph_before(appendix_anchor, "17.4 明确禁止的误判", "Heading 2")
    add_bullets_before(
        appendix_anchor,
        [
            "导航中存在页面或 Dashboard 出现卡片，不代表对应业务模块已经实现。",
            "排班记录不等于实际考勤，也不能直接作为已发生工资。",
            "配方预估价格字段存在，不代表采购价和产品成本闭环已经完成。",
            "开发环境能通过局域网访问，不代表已经完成公网生产部署。",
            "原规划中的 V1.0 必须项不等于当前已经完成；以第 13 章状态矩阵为当前快照。",
        ],
    )

    insert_paragraph_before(appendix_anchor, "18. 未来展望", "Heading 1")
    insert_paragraph_before(
        appendix_anchor,
        "BakeOps 的长期方向是把面包店每天分散的人员、配方、库存、生产、销售、损耗和活动记录，逐步连接成可解释、可追踪、可执行的经营系统。近期重点不是追求大量图表或 AI 功能，而是保证事实数据、库存流水和成本口径可靠。",
    )
    future_rows = [
        ["阶段", "目标", "预期价值"],
        ["业务闭环", "采购 → 库存 → 配方 → 计划 → 生产 → 销售/损耗 → 利润", "老板和店长能解释每天发生了什么，以及利润为何变化。"],
        ["主动运营", "安全库存、临期、生产缺口、高损耗和活动节点提醒", "从事后查看数据转向事前准备和异常处理。"],
        ["智能建议", "在拥有足够高质量历史数据后，结合星期、季节、节假日、营销与天气给出需求和产量建议", "降低缺货、过量生产和损耗；建议必须可解释并允许人工覆盖。"],
        ["组织扩展", "多门店、门店间比较、集中采购、跨店配方与权限", "支持业务从单店扩展，同时保留门店级数据隔离。"],
        ["系统化运营", "任务/SOP、文档、审计、通知、移动端快捷录入、外部 POS/订单接口", "减少重复手工工作并提高执行可追踪性。"],
    ]
    insert_table_before(document, appendix_anchor, future_rows, [3.1, 7.0, 8.7])

    insert_paragraph_before(appendix_anchor, "18.1 文档维护约定", "Heading 2")
    add_bullets_before(
        appendix_anchor,
        [
            "每次完成一个里程碑或改变核心架构/业务约定时提升文档小版本，并写明准确日期。",
            "实施状态必须由代码、迁移、测试和实际联调共同支持；页面草图或占位页只能标为部分实现。",
            "长期蓝图可以保持稳定；变化频繁的内容集中维护在第 13—18 章，便于后续工具快速定位。",
            "新增 ADR 记录不可逆或影响多个模块的架构决策，例如多门店、库存计价、销售接入和部署拓扑。",
        ],
    )

    for paragraph in document.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
        if paragraph.text.startswith("13.") or paragraph.text.startswith("14.") or paragraph.text.startswith("15.") or paragraph.text.startswith("16.") or paragraph.text.startswith("17.") or paragraph.text.startswith("18."):
            for run in paragraph.runs:
                if paragraph.style.name.startswith("Heading"):
                    run.font.color.rgb = RGBColor.from_string(ACCENT)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
